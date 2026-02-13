import streamlit as st
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import json
from datetime import datetime

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Pekhom Billing Manager", layout="wide", page_icon="🥘")

# --- COMPANY DETAILS ---
COMPANY_NAME = "PEKHOM FOOD SERVICE"
TAGLINE = "We deliver food for any kinds of parties, Office and other Orders."
PROP_NAME = "Sujoy Sil Sarma"
ADDRESS = "Chhanban, Udaipur, Gomati, Tripura."
GST_NO = "16FPSPS1850E1ZI"
BENEFICIARY_CODE = "1528516"
CONTACT = "8119914074"

BANK_DETAILS = {
    "Name": "SUJOY SIL SARMA",
    "A/c": "34751669032",
    "IFSC": "SBIN0016194",
    "Branch": "Ramesh chowmuni SBI"
}

HISTORY_FILE = "pekhom_history_v2.csv"

# --- 2. SESSION STATE ---
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'items_final_v3' not in st.session_state:
    st.session_state.items_final_v3 = []
if 'edit_mode_data' not in st.session_state:
    st.session_state.edit_mode_data = {}
if 'quick_desc' not in st.session_state:
    st.session_state.quick_desc = ""

# --- LOGIN SYSTEM ---
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center; margin-top: 50px;'>🥘 PEKHOM FOOD SERVICE</h1>", unsafe_allow_html=True)
    st.markdown("<h3 style='text-align: center; color: grey;'>Billing System Login</h3>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        st.write("")
        with st.form("login_form"):
            username = st.text_input("Login ID")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Login", use_container_width=True)
            
            if submitted:
                if username == "sujoyish2001" and password == "joyta2212":
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("❌ Incorrect ID or Password")
    st.stop() # Stops the rest of the app from loading until logged in

# --- 3. HELPER FUNCTIONS ---

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            df = pd.read_csv(HISTORY_FILE)
            if 'Items JSON' not in df.columns: df['Items JSON'] = "[]"
            if 'Payments' not in df.columns: df['Payments'] = 0.0
            if 'GST Rate' not in df.columns: df['GST Rate'] = 0
            if 'Client Addr' not in df.columns: df['Client Addr'] = ""
            if 'Client GST' not in df.columns: df['Client GST'] = ""
            return df
        except:
            pass
    return pd.DataFrame(columns=["Date", "Invoice No", "Client Name", "Total Amount", "Items JSON", "Payments", "GST Rate", "Client Addr", "Client GST"])

def save_to_history(date_str, inv_no, client, total, items, gst_rate, c_addr="", c_gst=""):
    df = load_history()
    if inv_no in df['Invoice No'].values:
        idx = df.index[df['Invoice No'] == inv_no][0]
        df.at[idx, 'Date'] = date_str
        df.at[idx, 'Client Name'] = client
        df.at[idx, 'Total Amount'] = total
        df.at[idx, 'Items JSON'] = json.dumps(items)
        df.at[idx, 'GST Rate'] = gst_rate
        df.at[idx, 'Client Addr'] = c_addr
        df.at[idx, 'Client GST'] = c_gst
    else:
        new_row = {
            "Date": date_str, "Invoice No": inv_no, "Client Name": client,
            "Total Amount": total, "Items JSON": json.dumps(items),
            "Payments": 0.0, "GST Rate": gst_rate,
            "Client Addr": c_addr, "Client GST": c_gst
        }
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    df.to_csv(HISTORY_FILE, index=False)

def update_payment(inv_no, amount):
    df = load_history()
    if inv_no in df['Invoice No'].values:
        idx = df.index[df['Invoice No'] == inv_no][0]
        current = df.at[idx, 'Payments']
        df.at[idx, 'Payments'] = current + amount
        df.to_csv(HISTORY_FILE, index=False)
        return True
    return False

def get_next_invoice_number():
    df = load_history()
    if df.empty:
        count = 1
    else:
        try:
            last_inv = df.iloc[-1]['Invoice No']
            parts = last_inv.split('-')
            count = int(parts[-1]) + 1
        except:
            count = len(df) + 1
    year = datetime.now().year
    return f"INV-{year}-{count:03d}"

# --- PDF GENERATOR (Black & White) ---
def generate_pdf(invoice_data, items_df, grand_total, sub_total, cgst_val, sgst_val):
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    
    # Border
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    pdf.rect(5, 5, 200, 287)
    
    pdf.set_font("Times", 'B', 24)
    pdf.set_text_color(0, 0, 0) 
    
    # Header
    pdf.set_y(15)
    pdf.cell(0, 10, COMPANY_NAME, ln=True, align='C')
    
    pdf.set_font("Times", 'I', 10)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 5, TAGLINE, ln=True, align='C')
    
    pdf.ln(5)
    y_start = pdf.get_y()
    
    # Left: Provider Info
    pdf.set_x(15)
    pdf.set_font("Times", 'B', 10)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(90, 5, f"Prop: {PROP_NAME}", ln=True)
    pdf.set_font("Times", '', 10)
    pdf.set_x(15); pdf.cell(90, 5, ADDRESS, ln=True)
    
    if not invoice_data.get('hide_gst', False):
        pdf.set_x(15); pdf.cell(90, 5, f"GSTIN: {GST_NO}", ln=True)
        
    pdf.set_x(15); pdf.cell(90, 5, f"Ben. Code: {BENEFICIARY_CODE}", ln=True)
    pdf.set_x(15); pdf.cell(90, 5, f"Contact: {CONTACT}", ln=True)
    
    # Right: Invoice Meta
    pdf.set_xy(110, y_start)
    pdf.set_font("Times", 'B', 14)
    pdf.cell(85, 8, "INVOICE", ln=True, align='R')
    pdf.set_font("Times", '', 10)
    pdf.set_x(110); pdf.cell(85, 5, f"Date: {invoice_data['display_date']}", ln=True, align='R')
    pdf.set_x(110); pdf.cell(85, 5, f"Invoice No: {invoice_data['invoice_no']}", ln=True, align='R')
    
    pdf.ln(8)
    pdf.set_draw_color(0, 0, 0)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)
    
    # Bill To
    pdf.set_x(15)
    pdf.set_font("Times", 'B', 11)
    pdf.cell(0, 5, "Bill To:", ln=True)
    pdf.set_font("Times", '', 10)
    pdf.set_x(15); pdf.cell(0, 5, invoice_data['client_name'], ln=True)
    if invoice_data.get('client_addr'):
        pdf.set_x(15); pdf.multi_cell(0, 5, invoice_data['client_addr'])
    if not invoice_data.get('hide_gst', False) and invoice_data.get('client_gst'):
        pdf.set_x(15); pdf.cell(0, 5, f"GSTIN: {invoice_data['client_gst']}", ln=True)
        
    pdf.ln(5)
    
    # Table Header
    pdf.set_x(10)
    pdf.set_fill_color(240, 240, 240)
    pdf.set_font("Times", 'B', 10)
    
    headers = ["Description"]
    if invoice_data.get('show_item_date', False): headers.append("Date")
    headers.extend(["Qty", "Price", "Total"])
    
    total_w = 190
    fixed_w = 20 + 30 + 35 
    desc_w = total_w - fixed_w
    if invoice_data.get('show_item_date', False):
        desc_w -= 25
        widths = [desc_w, 25, 20, 30, 35]
    else:
        widths = [desc_w, 20, 30, 35]
        
    for i, h in enumerate(headers):
        pdf.cell(widths[i], 8, h, 1, 0, 'C', True)
    pdf.ln()
    
    # Table Rows
    pdf.set_font("Times", '', 10)
    for item in items_df.to_dict('records'):
        pdf.set_x(10)
        # Support older records that might not have 'Total' calculated properly
        price = float(item.get('Price', 0))
        qty = float(item.get('Qty', 0))
        line_total = qty * price
        
        row = [item.get('Description', '')]
        if invoice_data.get('show_item_date', False):
            row.append(str(item.get('Item Date', '')))
        row.extend([str(item.get('Qty', '')), f"{price:.2f}", f"{line_total:.2f}"])
        for i, val in enumerate(row):
            align = 'L' if i == 0 else 'C'
            pdf.cell(widths[i], 8, str(val), 1, 0, align)
        pdf.ln()
        
    # Totals
    pdf.ln(5)
    def print_total_row(label, value, bold=False):
        pdf.set_font("Times", 'B' if bold else '', 10 if not bold else 12)
        pdf.set_x(10)
        pdf.cell(140, 6, label, 0, 0, 'R')
        pdf.cell(50, 6, value, 1 if bold else 0, 1, 'R')

    if not invoice_data.get('hide_gst', False) and invoice_data.get('gst_rate', 0) > 0:
        print_total_row("Sub Total:", f"{sub_total:.2f}")
        print_total_row(f"CGST ({invoice_data['gst_rate']/2}%):", f"{cgst_val:.2f}")
        print_total_row(f"SGST ({invoice_data['gst_rate']/2}%):", f"{sgst_val:.2f}")
    
    print_total_row("Grand Total:", f"Rs. {grand_total:.2f}", bold=True)
    
    pdf.ln(10)
    if pdf.get_y() > 240: pdf.add_page()
    y_footer = pdf.get_y()
    
    # Bank
    pdf.set_x(15)
    pdf.set_font("Times", 'B', 10)
    pdf.cell(100, 5, "Bank Account Details:", ln=True)
    pdf.set_font("Times", '', 9)
    pdf.set_x(15); pdf.cell(100, 5, f"Name: {BANK_DETAILS['Name']}", ln=True)
    pdf.set_x(15); pdf.cell(100, 5, f"A/c: {BANK_DETAILS['A/c']}", ln=True)
    pdf.set_x(15); pdf.cell(100, 5, f"IFSC: {BANK_DETAILS['IFSC']}", ln=True)
    pdf.set_x(15); pdf.cell(100, 5, f"Branch: {BANK_DETAILS['Branch']}", ln=True)
    
    # Signatory
    pdf.set_xy(120, y_footer + 5)
    pdf.cell(70, 5, "For PEKHOM FOOD SERVICE", ln=True, align='C')
    pdf.ln(8)
    pdf.set_x(120)
    pdf.cell(70, 5, "Authorized Signatory", ln=True, align='C')
    
    return pdf.output(dest='S').encode('latin-1')

# --- RECEIPT GENERATOR (B&W) ---
def generate_receipt_pdf(inv_no, client_name, date_val, amount_paid, total_due):
    pdf = FPDF(orientation='P', unit='mm', format='A5')
    pdf.add_page()
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    pdf.rect(5, 5, 138, 200) 
    
    pdf.set_font("Times", 'B', 18)
    pdf.set_text_color(0, 0, 0) 
    pdf.cell(0, 10, COMPANY_NAME, ln=True, align='C')
    
    pdf.set_font("Times", 'B', 14)
    pdf.ln(5)
    pdf.cell(0, 10, "PAYMENT RECEIPT", ln=True, align='C', border=1)
    
    pdf.ln(10)
    pdf.set_font("Times", '', 12)
    pdf.set_x(15); pdf.cell(40, 10, "Date:", 0, 0); pdf.cell(0, 10, str(date_val), 0, 1)
    pdf.set_x(15); pdf.cell(40, 10, "Receipt For:", 0, 0); pdf.cell(0, 10, f"Invoice No {inv_no}", 0, 1)
    pdf.set_x(15); pdf.cell(40, 10, "Received From:", 0, 0); pdf.cell(0, 10, client_name, 0, 1)
    
    pdf.ln(5)
    pdf.set_x(15)
    pdf.set_font("Times", 'B', 14)
    pdf.cell(60, 12, "Amount Received:", 1, 0, 'L')
    pdf.cell(50, 12, f"Rs. {amount_paid:.2f}", 1, 1, 'R')
    
    pdf.set_x(15)
    pdf.set_font("Times", '', 11)
    pdf.cell(60, 8, "Balance Due:", 1, 0, 'L')
    pdf.cell(50, 8, f"Rs. {total_due - amount_paid:.2f}", 1, 1, 'R')
    
    pdf.ln(20)
    pdf.set_x(80)
    pdf.cell(50, 5, "Authorized Signatory", 'T', 1, 'C')
    return pdf.output(dest='S').encode('latin-1')

# --- WORD GENERATOR ---
def generate_word(invoice_data, items_df, grand_total, sub_total, cgst_val, sgst_val):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    
    t = doc.add_heading(COMPANY_NAME, 0)
    t.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    t.runs[0].font.name = 'Times New Roman'
    doc.add_paragraph(TAGLINE).italic = True
    
    provider_info = f"Prop: {PROP_NAME}\nAddress: {ADDRESS}\n"
    if not invoice_data.get('hide_gst', False):
        provider_info += f"GSTIN: {GST_NO}\n"
    provider_info += f"Ben. Code: {BENEFICIARY_CODE}\nContact: {CONTACT}"
    doc.add_paragraph(provider_info)
    
    doc.add_heading("INVOICE", level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Date: {invoice_data['display_date']}\nInvoice No: {invoice_data['invoice_no']}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    client_info = f"Bill To: {invoice_data['client_name']}\n{invoice_data.get('client_addr', '')}"
    if not invoice_data.get('hide_gst', False) and invoice_data.get('client_gst'):
        client_info += f"\nGSTIN: {invoice_data['client_gst']}"
    doc.add_paragraph(client_info)
    
    cols = 5 if invoice_data.get('show_item_date', False) else 4
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ["Description"]
    if invoice_data.get('show_item_date', False): headers.append("Date")
    headers.extend(["Qty", "Price", "Total"])
    for i, h in enumerate(headers): hdr[i].text = h
        
    for item in items_df.to_dict('records'):
        row = table.add_row().cells
        line_total = float(item.get('Qty', 0)) * float(item.get('Price', 0))
        idx = 0
        row[idx].text = item.get('Description', ''); idx += 1
        if invoice_data.get('show_item_date', False): row[idx].text = str(item.get('Item Date', '')); idx += 1
        row[idx].text = str(item.get('Qty', '')); idx += 1
        row[idx].text = f"{float(item.get('Price', 0)):.2f}"; idx += 1
        row[idx].text = f"{line_total:.2f}"
            
    doc.add_paragraph()
    if not invoice_data.get('hide_gst', False) and invoice_data.get('gst_rate', 0) > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Sub Total: {sub_total:.2f}\nCGST: {cgst_val:.2f}\nSGST: {sgst_val:.2f}\n")
    
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    t.add_run(f"Grand Total: Rs. {grand_total:.2f}").bold = True
    
    doc.add_paragraph(f"\nBank Details:\nName: {BANK_DETAILS['Name']}\nA/c: {BANK_DETAILS['A/c']}\nIFSC: {BANK_DETAILS['IFSC']}\nBranch: {BANK_DETAILS['Branch']}")
    doc.add_paragraph("\n\nFor PEKHOM FOOD SERVICE\n\nAuthorized Signatory").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. MAIN APP ---

# Logout Button at Top Right
col_title, col_logout = st.columns([5, 1])
with col_title:
    st.write("")
with col_logout:
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.logged_in = False
        st.rerun()

tab1, tab2 = st.tabs(["📝 Create / Edit Invoice", "📊 History & Payments"])

# === TAB 1: CREATE / EDIT ===
with tab1:
    st.title("Invoice Manager")
    
    default_inv_no = get_next_invoice_number()
    default_client = ""
    default_addr = ""
    default_cgst = ""
    default_gst_rate_idx = 0
    
    if st.session_state.edit_mode_data:
        d = st.session_state.edit_mode_data
        default_inv_no = d['Invoice No']
        default_client = d['Client Name']
        default_addr = d.get('Client Addr', '')
        default_cgst = d.get('Client GST', '')
        default_gst_rate_idx = [0, 5, 12, 18, 28].index(d['GST Rate']) if d['GST Rate'] in [0, 5, 12, 18, 28] else 0
        st.info(f"✏️ Editing Invoice: {default_inv_no}")
        if st.button("Cancel Edit & Clear"):
            st.session_state.edit_mode_data = {}
            st.session_state.items_final_v3 = []
            st.rerun()

    c_s1, c_s2 = st.columns(2)
    with c_s1:
        inv_no = st.text_input("Invoice Number", value=default_inv_no)
        col_d1, col_d2 = st.columns([1, 1])
        with col_d1: inv_date = st.date_input("Date", datetime.now())
        with col_d2: 
            st.write(""); st.write("")
            manual_date = st.checkbox("Leave Date Blank (Pen)")
    with c_s2:
        hide_gst = st.checkbox("Hide GST Completely", value=False)
        if not hide_gst:
            gst_rate = st.selectbox("GST Rate (Common)", [0, 5, 12, 18, 28], index=default_gst_rate_idx)
        else:
            gst_rate = 0
            st.write("")
        show_item_date = st.checkbox("Show Date for each Item?", value=False)
        
    st.markdown("---")
    c1, c2 = st.columns(2)
    with c1:
        c_name = st.text_input("Client Name", value=default_client, placeholder="Enter Name")
        c_addr = st.text_area("Address", value=default_addr, height=70, placeholder="Enter Address")
    with c2:
        if not hide_gst:
            c_gst = st.text_input("Client GSTIN (Optional)", value=default_cgst)
        else:
            c_gst = ""

    st.markdown("---")
    st.subheader("Items")

    # QUICK ADD
    st.write("**Quick-Add Descriptions:**")
    q1, q2, q3, q4 = st.columns(4)
    if q1.button("Veg Meal"): st.session_state.quick_desc = "Veg Meal"
    if q2.button("Chicken Meal"): st.session_state.quick_desc = "Chicken Meal"
    if q3.button("Fish Meal"): st.session_state.quick_desc = "Fish Meal"
    if q4.button("Mutton Meal"): st.session_state.quick_desc = "Mutton Meal"

    with st.form("entry_form", clear_on_submit=True):
        if show_item_date: cols = st.columns([3, 2, 1, 1])
        else: cols = st.columns([4, 1, 1])
        
        with cols[0]: i_desc = st.text_input("Description", value=st.session_state.quick_desc)
        
        idx = 1
        i_date_val = None
        if show_item_date:
            with cols[idx]: i_date_val = st.date_input("Item Date")
            idx += 1
        with cols[idx]: i_qty = st.number_input("Qty", min_value=1, value=1)
        idx += 1
        with cols[idx]: i_price = st.number_input("Price", min_value=0.0, value=0.0)
            
        if st.form_submit_button("➕ Add Item") and i_desc:
            st.session_state.items_final_v3.append({
                "Description": i_desc, 
                "Item Date": str(i_date_val) if i_date_val else "",
                "Qty": i_qty, 
                "Price": i_price
            })
            st.session_state.quick_desc = ""
            st.rerun()

    if st.session_state.items_final_v3:
        df = pd.DataFrame(st.session_state.items_final_v3)
        df['Total'] = df['Qty'] * df['Price']
        disp_cols = ["Description"]
        if show_item_date: disp_cols.append("Item Date")
        disp_cols.extend(["Qty", "Price", "Total"])
        st.dataframe(df[disp_cols], use_container_width=True)
        
        if st.button("❌ Remove Last"):
            st.session_state.items_final_v3.pop(); st.rerun()
            
        sub_total = df['Total'].sum()
        cgst_val = 0; sgst_val = 0
        if not hide_gst and gst_rate > 0:
            cgst_val = sub_total * (gst_rate / 2 / 100)
            sgst_val = sub_total * (gst_rate / 2 / 100)
        grand_total = sub_total + cgst_val + sgst_val
        
        st.markdown(f"<div style='text-align:right; padding:10px; background:#f0f2f6;'><h3 style='color:#000000;'>Total: Rs. {grand_total:.2f}</h3></div>", unsafe_allow_html=True)
        st.markdown("---")
        
        if st.button("✅ GENERATE / UPDATE BILL", type="primary", use_container_width=True):
            if not c_name: st.error("Name Required!")
            else:
                final_date_str = "____________________" if manual_date else str(inv_date)
                inv_data = {
                    'client_name': c_name, 'client_addr': c_addr, 'client_gst': c_gst,
                    'invoice_no': inv_no, 'display_date': final_date_str, 
                    'show_item_date': show_item_date, 'gst_rate': gst_rate,
                    'hide_gst': hide_gst
                }
                pdf_data = generate_pdf(inv_data, df, grand_total, sub_total, cgst_val, sgst_val)
                word_data = generate_word(inv_data, df, grand_total, sub_total, cgst_val, sgst_val)
                
                # Save including address and GST
                save_to_history(final_date_str, inv_no, c_name, grand_total, st.session_state.items_final_v3, gst_rate, c_addr, c_gst)
                
                st.success("Success!")
                safe_name = "".join(x for x in c_name if x.isalnum() or x in " _-")
                f_name = f"{safe_name}_{inv_no}"
                d1, d2 = st.columns(2)
                with d1: st.download_button("📄 PDF (Black & White)", pdf_data, f"{f_name}.pdf", "application/pdf")
                with d2: st.download_button("📝 Word", word_data, f"{f_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# === TAB 2: HISTORY & PAYMENTS ===
with tab2:
    st.title("Ledger & Payments")
    hist_df = load_history()
    
    if not hist_df.empty:
        all_invoices = hist_df['Invoice No'].unique().tolist()
        selected_inv = st.selectbox("Select Invoice to Manage:", all_invoices)
        row = hist_df[hist_df['Invoice No'] == selected_inv].iloc[0]
        
        st.markdown(f"### Invoice: {row['Invoice No']} | {row['Client Name']}")
        col_st1, col_st2, col_st3 = st.columns(3)
        total_amt = float(row.get('Total Amount', 0))
        paid_amt = float(row.get('Payments', 0))
        due_amt = total_amt - paid_amt
        
        col_st1.metric("Total Bill", f"Rs. {total_amt:.2f}")
        col_st2.metric("Paid So Far", f"Rs. {paid_amt:.2f}")
        col_st3.metric("Balance Due", f"Rs. {due_amt:.2f}", delta_color="inverse" if due_amt > 0 else "normal")
        
        with st.expander("View Bill Items"):
            try:
                items_data = json.loads(row['Items JSON'])
                st.table(pd.DataFrame(items_data))
            except:
                st.warning("Item details unavailable for old records.")

        st.markdown("---")
        # 3 Columns now
        c_act1, c_act2, c_act3 = st.columns(3)
        
        with c_act1:
            st.subheader("💰 Add Payment")
            pay_amount = st.number_input("Amount Received", min_value=0.0, max_value=float(due_amt) if due_amt > 0 else 0.0, value=0.0)
            if st.button("Record Payment & Get Receipt", use_container_width=True):
                if pay_amount > 0:
                    update_payment(selected_inv, pay_amount)
                    rec_pdf = generate_receipt_pdf(selected_inv, row['Client Name'], datetime.now().date(), pay_amount, due_amt)
                    st.success("Payment Recorded!")
                    st.download_button("📥 Download Receipt", rec_pdf, f"Receipt_{selected_inv}.pdf", "application/pdf")
                    st.rerun()
                else:
                    st.error("Enter valid amount")

        with c_act2:
            st.subheader("✏️ Edit / Load")
            if st.button("Load Data into Editor", use_container_width=True):
                try:
                    items_data = json.loads(row['Items JSON'])
                    st.session_state.items_final_v3 = items_data
                    st.session_state.edit_mode_data = {
                        'Invoice No': row['Invoice No'],
                        'Client Name': row['Client Name'],
                        'GST Rate': row.get('GST Rate', 0),
                        'Client Addr': row.get('Client Addr', ''),
                        'Client GST': row.get('Client GST', '')
                    }
                    st.success("Loaded! Go to 'Create Invoice' tab.")
                except:
                    st.error("Cannot load old data format.")
                    
        with c_act3:
            st.subheader("📄 Download Old Bill")
            # Generate the PDF dynamically from the JSON
            try:
                old_items = json.loads(row['Items JSON'])
                if old_items:
                    df_old = pd.DataFrame(old_items)
                    has_date = 'Item Date' in df_old.columns
                    old_gst_r = row.get('GST Rate', 0)
                    
                    old_sub = sum(float(x.get('Qty',0)) * float(x.get('Price',0)) for x in old_items)
                    old_cgst = old_sub * (old_gst_r / 2 / 100) if old_gst_r > 0 else 0
                    old_sgst = old_sub * (old_gst_r / 2 / 100) if old_gst_r > 0 else 0
                    
                    old_data = {
                        'client_name': row['Client Name'],
                        'client_addr': row.get('Client Addr', ''),
                        'client_gst': row.get('Client GST', ''),
                        'invoice_no': row['Invoice No'],
                        'display_date': row['Date'],
                        'show_item_date': has_date,
                        'gst_rate': old_gst_r,
                        'hide_gst': old_gst_r == 0
                    }
                    old_pdf_bytes = generate_pdf(old_data, df_old, total_amt, old_sub, old_cgst, old_sgst)
                    st.download_button(
                        label="📥 Download Bill PDF",
                        data=old_pdf_bytes,
                        file_name=f"{row['Invoice No']}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
            except Exception as e:
                st.error("Data too old to regenerate directly. Use 'Edit' instead.")
                
        st.markdown("---")
        st.subheader("Full History")
        st.dataframe(hist_df.drop(columns=['Items JSON']), use_container_width=True)
    else:
        st.info("No history found.")